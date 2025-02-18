from docx import Document
from docx.shared import Pt, Inches
import json
import pandas as pd

def inspect_document(doc_path):
    """Inspect a Word document's content and formatting"""
    doc = Document(doc_path)
    
    document_data = {
        "paragraphs": [],
        "sections": [],
        "heading_styles": {}  # Add heading styles inspection
    }
    
    # Inspect heading styles
    for i in range(1, 10):  # Check Heading 1 through 9
        style_name = f'Heading {i}'
        if style_name in doc.styles:
            style = doc.styles[style_name]
            document_data["heading_styles"][style_name] = {
                "font_name": style.font.name,
                "font_size": style.font.size.pt if style.font.size else None,
                "bold": style.font.bold,
                "italic": style.font.italic
            }
    
    # Inspect paragraphs
    for i, paragraph in enumerate(doc.paragraphs):
        style_font = paragraph.style.font
        para_data = {
            "index": i,
            "text": paragraph.text,
            "style": paragraph.style.name,
            "alignment": str(paragraph.alignment),
            "style_font": {
                "name": style_font.name,
                "size": style_font.size.pt if style_font.size else None,
                "bold": style_font.bold,  # Add bold property
                "italic": style_font.italic  # Add italic property
            },
            "runs": []
        }
        
        # Inspect runs
        for j, run in enumerate(paragraph.runs):
            run_data = {
                "index": j,
                "text": run.text,
                "font": {
                    "name": run.font.name or "Inherited",
                    "size": run.font.size.pt if run.font.size else None,
                    "bold": run.bold,
                    "italic": run.italic
                }
            }
            para_data["runs"].append(run_data)
            
        document_data["paragraphs"].append(para_data)
    
    # Inspect sections
    for section in doc.sections:
        section_data = {
            "dimensions": {
                "page_height": f"{section.page_height.inches:.2f}",
                "page_width": f"{section.page_width.inches:.2f}",
                "left_margin": f"{section.left_margin.inches:.2f}",
                "right_margin": f"{section.right_margin.inches:.2f}"
            },
            "header": section.header.paragraphs[0].text if section.header else None,
            "footer": section.footer.paragraphs[0].text if section.footer else None
        }
        document_data["sections"].append(section_data)
    
    return document_data

def convert_to_dataframe(document_data):
    """Convert document inspection data to pandas DataFrames"""
    # Create DataFrame for paragraphs and runs
    paragraph_rows = []
    for para in document_data['paragraphs']:
        style_font = para['style_font']
        # Flatten paragraph data
        base_row = {
            'paragraph_index': para['index'],
            'text': para['text'],
            'style': para['style'],
            'alignment': para['alignment'],
            'font_name': style_font['name'],
            'font_size': style_font['size'],
            'font_bold': style_font['bold'],
            'font_italic': style_font['italic']
        }
        
        # Add run-level information
        for run in para['runs']:
            run_row = base_row.copy()
            run_row.update({
                'run_index': run['index'],
                'run_text': run['text'],
                'run_font_name': run['font']['name'],
                'run_font_size': run['font']['size'],
                'run_bold': run['font']['bold'],
                'run_italic': run['font']['italic']
            })
            paragraph_rows.append(run_row)
    
    # Create DataFrame for sections
    section_rows = []
    for i, section in enumerate(document_data['sections']):
        section_rows.append({
            'section_index': i,
            'page_height': float(section['dimensions']['page_height']),
            'page_width': float(section['dimensions']['page_width']),
            'left_margin': float(section['dimensions']['left_margin']),
            'right_margin': float(section['dimensions']['right_margin']),
            'header': section['header'],
            'footer': section['footer']
        })
    
    # Create DataFrame for heading styles
    heading_rows = []
    for style_name, style_data in document_data['heading_styles'].items():
        heading_rows.append({
            'style_name': style_name,
            'font_name': style_data['font_name'],
            'font_size': style_data['font_size'],
            'bold': style_data['bold'],
            'italic': style_data['italic']
        })
    
    return {
        'paragraphs': pd.DataFrame(paragraph_rows),
        'sections': pd.DataFrame(section_rows),
        'heading_styles': pd.DataFrame(heading_rows)
    }

if __name__ == "__main__":
    # Inspect document
    doc_data = inspect_document(r"generated_docs\20250113 DTP Business Plan (1).docx")
    
    # Convert to DataFrames
    dataframes = convert_to_dataframe(doc_data)
    
    # Save DataFrames to CSV files
    dataframes['paragraphs'].to_csv('doc_paragraphs.csv', index=False)
    dataframes['sections'].to_csv('doc_sections.csv', index=False)
    dataframes['heading_styles'].to_csv('doc_heading_styles.csv', index=False)
    
    # Display DataFrames
    print("\nParagraphs and Runs:")
    print(dataframes['paragraphs'])
    print("\nSections:")
    print(dataframes['sections'])
    print("\nHeading Styles:")
    print(dataframes['heading_styles'])
    # Print formatted JSON
    print(json.dumps(doc_data, indent=2))