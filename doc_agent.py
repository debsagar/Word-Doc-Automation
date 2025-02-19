from langchain.agents import Tool, AgentExecutor, LLMSingleActionAgent
from langchain.memory import ConversationBufferMemory
from langchain.chat_models import ChatOpenAI
from langchain.prompts import StringPromptTemplate
from langchain.chains import LLMChain
from docx import Document
import os
from typing import List, Union
import json

class DocumentAgent:
    def __init__(self, api_key: str):
        self.llm = ChatOpenAI(
            temperature=0,
            model_name="gpt-3.5-turbo",
            openai_api_key=api_key
        )
        
        # Define tools
        self.tools = [
            Tool(
                name="Read_Document",
                func=self.read_document,
                description="Read content from a Word document"
            ),
            Tool(
                name="Create_Document",
                func=self.create_document,
                description="Create a new Word document with specified content"
            ),
            Tool(
                name="Update_Document",
                func=self.update_document,
                description="Update content in an existing Word document"
            )
        ]
        
        self.memory = ConversationBufferMemory(memory_key="chat_history")
    
    def read_document(self, file_path: str) -> str:
        """Read content from a Word document"""
        try:
            doc = Document(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return content
        except Exception as e:
            return f"Error reading document: {str(e)}"
    
    def create_document(self, content: Union[str, dict]) -> str:
        """Create a new Word document with specified content"""
        try:
            doc = Document()
            
            # If content is a JSON string, parse it
            if isinstance(content, str):
                try:
                    content = json.loads(content)
                except:
                    doc.add_paragraph(content)
                    
            # If content is a dictionary, format it properly
            if isinstance(content, dict):
                for key, value in content.items():
                    doc.add_heading(key, level=1)
                    doc.add_paragraph(str(value))
            
            # Save document
            os.makedirs('generated_docs', exist_ok=True)
            filename = f'generated_doc_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            filepath = os.path.join('generated_docs', filename)
            doc.save(filepath)
            return f"Document created successfully: {filepath}"
            
        except Exception as e:
            return f"Error creating document: {str(e)}"
    
    def update_document(self, file_path: str, updates: Union[str, dict]) -> str:
        """Update content in an existing Word document"""
        try:
            doc = Document(file_path)
            
            # If updates is a JSON string, parse it
            if isinstance(updates, str):
                try:
                    updates = json.loads(updates)
                except:
                    doc.add_paragraph(updates)
            
            # If updates is a dictionary, apply updates
            if isinstance(updates, dict):
                for key, value in updates.items():
                    doc.add_heading(key, level=1)
                    doc.add_paragraph(str(value))
            
            doc.save(file_path)
            return f"Document updated successfully: {file_path}"
            
        except Exception as e:
            return f"Error updating document: {str(e)}"

# Example usage
if __name__ == "__main__":
    # Initialize agent with your OpenAI API key
    agent = DocumentAgent("your-api-key-here")
    
    # Example: Create a document
    content = {
        "Title": "Meeting Minutes",
        "Summary": "Discussion about project progress",
        "Action Items": "1. Complete documentation\n2. Schedule follow-up"
    }
    
    result = agent.create_document(content)
    print(result)
    
    # Example: Read a document
    doc_content = agent.read_document("generated_docs/sample.docx")
    print(doc_content) 