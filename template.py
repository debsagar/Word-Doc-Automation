from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_template():
    """Create a template document for meeting minutes"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Add header
    header_section = doc.sections[0]
    header = header_section.header
    header_text = header.paragraphs[0]
    header_text.text = "Meeting Minutes"
    header_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Meeting Title and Date
    doc.add_heading('Meeting Title and Date: {}', level=0)
    
    # Meeting Summary
    doc.add_heading('Meeting Summary', level=1)
    doc.add_paragraph('{}')
    
    # Key Discussions & Findings
    doc.add_heading('Key Discussions & Findings', level=1)
    doc.add_paragraph('{}')
    
    # Decisions and Action Items
    doc.add_heading('Decisions and Action Items', level=1)
    doc.add_paragraph('{}')
    
    # Team Assignments
    doc.add_heading('Team Assignments', level=1)
    doc.add_paragraph('{}')
    
    # Next Meeting
    doc.add_heading('Next Meeting', level=1)
    doc.add_paragraph('{}')
    
    # Footer with generation date
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Generated on: {}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Save template
    doc.save('template.docx')
    print("Meeting minutes template created successfully!")

if __name__ == "__main__":
    create_template()