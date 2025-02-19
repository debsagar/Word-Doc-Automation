from docx import Document
from docx.shared import Inches
import os

# Create a directory for saving documents if it doesn't exist
output_dir = 'image_docs'
os.makedirs(output_dir, exist_ok=True)

# Create a new Word document
document = Document()

# Add empty paragraphs to push the image to the bottom
for _ in range(10):  # Adjust this number to move image up/down
    document.add_paragraph('')

# Add an image
# Replace 'test.jpeg' with your actual image path
image = document.add_picture('test.jpeg', width=Inches(3))

# Get the last paragraph (which contains the image)
last_paragraph = document.paragraphs[-1]
# Center align the image
last_paragraph.alignment = 1  # 1 represents center alignment

# Save the document in the output directory
document.save(os.path.join(output_dir, 'document_with_image1.docx'))