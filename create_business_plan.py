from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_cover_page():
    doc = Document()
    
    # Set margins
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Add logo
    doc.add_picture('test.jpeg', width=Inches(1.5))
    
    # Add some vertical space
    for _ in range(3):
        doc.add_paragraph()
    
    # Add subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_run = subtitle.add_run("Your Partner in AI-Transformation")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
    
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("Digital Technology Partner\n")
    title_run.font.size = Pt(24)
    title_run.bold = True
    
    # Add subtitle (Business Plan)
    plan_title = title.add_run("Business Plan")
    plan_title.font.size = Pt(24)
    plan_title.font.color.rgb = RGBColor(0, 255, 255)  # Cyan
    
    # Add date
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date.add_run("12th January 2025")
    
    # Add website at the bottom
    for _ in range(3):
        doc.add_paragraph()
    website = doc.add_paragraph()
    website.alignment = WD_ALIGN_PARAGRAPH.LEFT
    website_run = website.add_run("www.digitaltechnologypartner.ai")
    website_run.font.size = Pt(8)
    website_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save the document
    os.makedirs('generated_docs', exist_ok=True)
    doc.save('generated_docs/cover_page.docx')

if __name__ == "__main__":
    create_cover_page()
