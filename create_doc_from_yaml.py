from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
import yaml
import lorem  # for generating dummy text
import os

def add_styled_text(paragraph, text, size=None, bold=False, color=None, alignment=None):
    run = paragraph.add_run(text)
    font = run.font
    if size:
        font.size = Pt(size)
    run.bold = bold
    if color:
        font.color.rgb = RGBColor(*color)
    if alignment:
        paragraph.alignment = alignment
    return run

def create_document_from_yaml(yaml_path):
    # Read YAML file
    with open(yaml_path, 'r') as file:
        structure = yaml.safe_load(file)
    
    # Create document
    doc = Document()
    
    # Add logo
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    run.add_picture(structure['company_info']['logo'], width=Inches(0.8))  # Reduced size
    
    # Add space
    doc.add_paragraph()
    
    # Add company name
    add_styled_text(doc.add_paragraph(), structure['company_info']['name'], 
                   size=24, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Add document title
    add_styled_text(doc.add_paragraph(), structure['company_info']['document_title'],
                   size=20, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Add subtitle
    add_styled_text(doc.add_paragraph(), structure['company_info']['document_subtitle'],
                   size=16, color=(100, 100, 100), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Add space
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add prepared for/by section
    details_para = doc.add_paragraph()
    details_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_styled_text(details_para, f"Prepared for:\n", size=11, bold=True)
    add_styled_text(details_para, f"{structure['company_info']['prepared_for']}\n\n", size=11)
    
    add_styled_text(details_para, "Prepared by:\n", size=11, bold=True)
    for author in structure['company_info']['prepared_by']:
        add_styled_text(details_para, f"{author}\n", size=11)
    
    # Add space
    doc.add_paragraph()
    
    # Add date and version
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_styled_text(footer_para, f"Date: {structure['company_info']['document_date']}\n", size=11)
    add_styled_text(footer_para, f"Version: {structure['company_info']['version']}\n", size=11)
    add_styled_text(footer_para, structure['company_info']['confidentiality'], 
                   size=11, bold=True, color=(192, 0, 0))  # Dark red for confidentiality
    
    # Add page break
    doc.add_page_break()
    
    # Add sections with dummy content
    for section in structure['sections']:
        # Add section heading
        heading = doc.add_heading(section['title'], level=section['level'])
        
        # Add dummy content
        para = doc.add_paragraph(lorem.paragraph())
        
        # Add some space between sections
        doc.add_paragraph()
    
    # Save the document
    output_dir = 'generated_docs'
    os.makedirs(output_dir, exist_ok=True)
    doc.save(os.path.join(output_dir, 'document_from_yaml.docx'))

if __name__ == "__main__":
    create_document_from_yaml('document_structure.yaml') 